# Created to aid in the training of our NN.
# This program creates our character training samples in MS Docs.
# The program requires an input of a csv file with all the font names that you will want to be included in the MS doc.

from docx import Document
from docx.shared import Pt
import csv

# Removes duplicate fonts
with open('fontList.csv', encoding='utf-8', errors='ignore') as fontList, open('fontListDeDup.csv', 'w') as out_file:
    seen = set()  # set for fast O(1) amortized lookup
    reader = csv.DictReader(fontList)

    for line in fontList:  # duplication removal
        if line in seen: continue  # skip duplicate

        seen.add(line)
        out_file.write(line)

# loads all fonts into a dictionary object
with open('fontListDeDup.csv', encoding='utf-8', errors='ignore') as deDupfontList:
    reader = csv.DictReader(deDupfontList)
    fontDict = {}
    for row in reader:  # csv to dict{}
        for header, value in row.items():
            try:
                fontDict[header].append((value.lstrip().rstrip()))
            except KeyError:
                fontDict[header] = [value]

charDoc = Document()

style = charDoc.styles['Normal']
font = style.font
font.size = Pt(24)

paragraph = charDoc.add_paragraph()

i = 0
j = 0
t = 0

for one in range(33, 127):
    print(one, ': ', chr(one))
    if j == 0:
        paragraph.add_run(chr(one))
        t +=1
        print('total', j)
        j += 1
    else:
        paragraph.add_run("    ")
        newChar = paragraph.add_run(chr(one))
        t += 1
        print('total', j)
        j += 1

    for font in fontDict['Name']:
        paragraph.add_run("    ")
        newChar = paragraph.add_run(chr(one))
        t += 1
        print('UTF-8: ', one, ': ', chr(one))
        print(i, ' font', font)
        newChar.font.name = font

        i += 1
        print('Total Characters = ', t)


charDoc.save("CharTrainingSet.docx")